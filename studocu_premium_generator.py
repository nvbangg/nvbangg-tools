"""
Tool tạo file tài liệu ngẫu nhiên để upload lên Studocu lấy Premium.
Để dùng Nhớ cài thư viện: pip install python-docx
"""

import os, random, uuid
from datetime import datetime
from docx import Document  # pip install python-docx

# === CẤU HÌNH ===
SO_LUONG_FILE_TAO = 10
THU_MUC = "studocu_output"
TU_VUNG = """
tài liệu nội dung hệ thống dữ liệu phân tích tổng hợp phương pháp kiểm tra hướng dẫn kết quả thông tin
nghiên cứu báo cáo đánh giá thực hiện triển khai xử lý quản lý cập nhật thiết kế xây dựng phát triển
giải pháp ứng dụng công nghệ chất lượng hiệu quả tối ưu chiến lược mục tiêu kế hoạch dự án
khảo sát thống kê mô hình cấu trúc quy trình tiêu chuẩn yêu cầu đặc điểm tính năng chức năng
bảo mật an toàn tích hợp liên kết giao diện trải nghiệm người dùng khách hàng đối tác
tài nguyên nguồn lực ngân sách chi phí lợi ích rủi ro cơ hội thách thức xu hướng thị trường
đào tạo hỗ trợ tư vấn chuyên gia kiến thức kỹ năng kinh nghiệm sáng tạo đổi mới cải tiến
""".split()

# Xóa thư mục cũ và tạo mới
try:
    os.makedirs(THU_MUC, exist_ok=True)
    for f in os.listdir(THU_MUC):
        os.remove(os.path.join(THU_MUC, f))
except Exception:
    print("❌ Lỗi không xóa/tạo được thư mục")
    exit(1)


# Hàm tạo nội dung ngẫu nhiên
def tao_cau():
    return " ".join(
        random.choices(TU_VUNG, k=random.randint(5, 20))
    ).capitalize() + random.choice([".", "!", "?", "...", ";", ":"])


def tao_doan():
    return " ".join(tao_cau() for _ in range(random.randint(5, 10)))


# Tạo các file
so_file_da_tao = 0
for i in range(SO_LUONG_FILE_TAO):
    try:
        id_doc = str(uuid.uuid4())
        doc = Document()
        doc.add_heading(f"Tài liệu #{i + 1} {id_doc}", level=1)
        doc.add_paragraph(f"Tạo lúc: {datetime.now()}\nSeed: {random.getstate()[1][0]}")
        for _ in range(random.randint(5, 15)):
            doc.add_paragraph(tao_doan())
        doc.add_paragraph("Các điểm chính:")
        for _ in range(random.randint(5, 10)):
            doc.add_paragraph(tao_cau(), style="List Bullet")
        ten_file = f"docx_{i + 1}_{id_doc}.docx"
        doc.save(os.path.join(THU_MUC, ten_file))
        print(f"✓ Đã tạo: {ten_file}")
        so_file_da_tao += 1
    except Exception:
        print(f"❌ Lỗi không thể tạo file #{i + 1}")

print(f"\n✅ Hoàn tất! Đã tạo {so_file_da_tao} file trong thư mục '{THU_MUC}'")
